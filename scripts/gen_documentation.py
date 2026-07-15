"""Generate AdhiKaar_Documentation.docx - full technical + product documentation."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PRIMARY = RGBColor(0x43, 0x38, 0xCA)
DARK = RGBColor(0x1C, 0x19, 0x17)
GREY = RGBColor(0x57, 0x53, 0x4E)

doc = Document()

# Base style
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(11)
st.font.color.rgb = DARK

def h1(t):
    p = doc.add_heading(level=1); r = p.add_run(t); r.font.color.rgb = PRIMARY; r.font.size = Pt(18); return p
def h2(t):
    p = doc.add_heading(level=2); r = p.add_run(t); r.font.color.rgb = PRIMARY; r.font.size = Pt(14); return p
def h3(t):
    p = doc.add_heading(level=3); r = p.add_run(t); r.font.color.rgb = RGBColor(0x31,0x2E,0x81); r.font.size = Pt(12); return p
def body(t):
    p = doc.add_paragraph(t); p.paragraph_format.space_after = Pt(6); return p
def bullet(t):
    p = doc.add_paragraph(t, style='List Bullet'); return p
def num(t):
    p = doc.add_paragraph(t, style='List Number'); return p
def kv_table(rows, widths=(2.2, 4.0)):
    tbl = doc.add_table(rows=0, cols=2); tbl.style = 'Light Grid Accent 1'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        c = tbl.add_row().cells
        c[0].paragraphs[0].add_run(k).bold = True
        c[1].text = v
    for r in tbl.rows:
        r.cells[0].width = Inches(widths[0]); r.cells[1].width = Inches(widths[1])
    doc.add_paragraph()
    return tbl

# ── Cover ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('अधिKaar'); r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = PRIMARY
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('An Offline-First, Verified, Multilingual AI Legal Assistant for Every Indian Citizen')
r.font.size = Pt(14); r.font.color.rgb = GREY
tag = doc.add_paragraph(); tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run('Your Rights, Your Language'); r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph(); doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Technical & Product Documentation\nBuilt with Gemma · Runs 100% on-device').font.color.rgb = GREY
doc.add_page_break()

# ── 1. Executive Summary ──
h1('1. Executive Summary')
body('अधिKaar ("Adhikaar", meaning "rights") is a legal-navigation assistant that helps ordinary Indian '
     'citizens understand their legal rights, decode legal documents, and take concrete next steps - in 11 '
     'Indian languages, by voice or text, and entirely on their own device. Unlike cloud chatbots, every '
     'component runs locally: the language model, the legal knowledge base, speech recognition, speech synthesis '
     'and document OCR. Nothing a user types, says, or uploads ever leaves their computer.')
body('The system is powered by Google\'s Gemma model running locally through Ollama, grounded on 6,845 reviewed '
     'chunks of official Indian law via retrieval-augmented generation (RAG), and hardened with a claim-level '
     'verification pipeline so that legal citations are checked against real statute text rather than model memory.')

# ── 2. Problem Statement ──
h1('2. Problem Statement')
h2('2.1 The Problem')
body('Over a billion Indians face everyday legal situations - unpaid wages, withheld security deposits, FIRs, '
     'domestic violence, consumer disputes - yet meaningful legal help remains out of reach. Lawyers are '
     'unaffordable for most; official information is English-first, scattered, and written in dense legalese; and '
     'the recent overhaul of India\'s criminal laws (IPC replaced by the Bharatiya Nyaya Sanhita, CrPC by the '
     'Bharatiya Nagarik Suraksha Sanhita, both effective 1 July 2024) has made even existing guidance outdated.')
h2('2.2 Target Audience')
bullet('Citizens who cannot afford a lawyer and do not know their rights.')
bullet('Non-English speakers, served in 11 Indian languages by voice and text.')
bullet('Low-literacy and elderly users, served through voice input, read-aloud output, and plain-language explanations.')
bullet('Anyone handling a legal document (FIR, notice, summons) they cannot understand.')
h2('2.3 Why It Matters')
body('Not knowing your rights often means losing them - a missed limitation deadline, an unfiled complaint, an '
     'unchallenged illegal eviction. And because legal matters are deeply personal, cloud-based AI raises real '
     'privacy fears. अधिKaar addresses both: it democratises legal knowledge and keeps it private by design.')

# ── 3. Solution Overview ──
h1('3. Solution Overview')
body('अधिKaar is a single-page web application backed by a local Python service. A user describes their situation '
     'in their own language (typed or spoken); the assistant confirms understanding, retrieves date-applicable '
     'official law, produces a verified, plain-language answer, and offers concrete tools to act on it.')
h2('3.1 Feature Suite')
kv_table([
    ('Talk to Legal Helper', 'Conversational, RAG-grounded legal guidance with dynamic response depth, power-imbalance detection, and multilingual voice in/out.'),
    ('Law & Next Steps', 'One verified analysis in six panels: situation & applicable law, per-claim verification, official sources with links, both-sides stress test, a shareable rights card, and a plain "explain to someone you trust" summary.'),
    ('Section Converter', 'Instant IPC↔BNS and CrPC↔BNSS conversion (216 + 80 curated mappings) with subsection-insensitive matching and offence-name search.'),
    ('Translate Legal Document', 'Upload a photo/PDF of a notice or FIR; on-device OCR extracts the text, the model explains it in plain words, and the user can ask questions answered strictly from that document.'),
    ('Draft a Document', '45+ ready-to-file citizen documents (FIR requests, RTI appeals, legal notices, rent agreements, wills, POAs, complaints) generated in the user\'s language.'),
    ('Find Legal Aid', 'PAN-India legal-aid directory: State Legal Services Authorities for all 28 states + 8 UTs, Delhi district DLSAs, and national helplines (NALSA 15100, Tele-Law 14454, etc.).'),
    ('Evidence Checklists', '20 situation-specific checklists of documents, steps and statutory deadlines, matched deterministically to the user\'s case.'),
    ('Rights Card / Elder Mode / Consequence Simulator', 'Shareable sourced rights card, an intermediary-friendly explanation, and a "what if you do nothing" deterministic view.'),
])

# ── 4. Technical Architecture ──
h1('4. Technical Architecture')
h2('4.1 High-Level Design')
body('The application follows a thin-client / local-service architecture. A zero-build vanilla-JavaScript '
     'single-page app runs in the browser; a Flask service on 127.0.0.1:5000 serves it and exposes a REST API. '
     'All heavy computation - LLM inference, embedding, retrieval, OCR, speech - happens inside the local service '
     'or, for OCR/voice-in-browser, on the client itself. No external API is called at runtime.')
body('Data flow: user input → (optional) multilingual query expansion → RAG retrieval from ChromaDB → prompt '
     'assembly with retrieved official-law context → local Gemma generation via Ollama → (for verified answers) '
     'claim-level verification → response rendered in the SPA.')
h2('4.2 Technology Stack')
kv_table([
    ('Frontend', 'Vanilla HTML/CSS/JavaScript single-page app, no build step. Vendored offline assets (Lucide icons, Marked, pdf.js, Tesseract.js, Noto fonts).'),
    ('Backend', 'Python 3 + Flask + Flask-CORS, loopback-only (127.0.0.1).'),
    ('LLM runtime', 'Ollama serving Gemma locally (gemma4:e4b primary, gemma3:4b fallback).'),
    ('Vector store / RAG', 'ChromaDB (persistent) with the default sentence-transformers embedding function (all-MiniLM-L6-v2); EmbeddingGemma available for semantic retrieval.'),
    ('Speech-to-text', 'faster-whisper ("small", int8, CPU) with VAD filtering and hallucination guards.'),
    ('Text-to-speech', 'Facebook MMS-TTS (per-language VITS models via transformers) with browser Web Speech API fallback.'),
    ('Document OCR', 'PaddleOCR (PP-OCRv6/PP-OCRv5, Devanagari + Latin) server-side; pdf.js text-layer extraction + Tesseract.js OCR fallback in-browser.'),
    ('PDF handling', 'pypdf (text layer) and pypdfium2 (rasterisation, no poppler needed).'),
])
h2('4.3 Backend API Endpoints')
kv_table([
    ('/api/chat', 'Main conversational endpoint with RAG grounding, keyword expansion, power-imbalance detection and per-session uploaded-document context.'),
    ('/api/law-and-steps', 'Verified six-panel analysis via the draft→verify→repair pipeline.'),
    ('/api/bns-convert, /api/crpc-convert', 'IPC↔BNS and CrPC↔BNSS section conversion.'),
    ('/api/extract-document, /api/upload-document, /api/ask-document, /api/clear-document', 'OCR extraction, per-session indexing, grounded document Q&A, and cleanup.'),
    ('/api/translate-document', 'Plain-language explanation of an extracted legal document.'),
    ('/api/draft-document, /api/document-templates', 'Document generation and the template catalogue.'),
    ('/api/legal-aid, /api/evidence-checklists', 'Legal-aid directory and evidence checklists.'),
    ('/api/rights-card, /api/rights-checklist, /api/consequence-simulator, /api/panchayat-bridge, /api/devil-advocate', 'Rights card, checklist, consequences, elder explanation and adversarial stress test.'),
    ('/api/transcribe, /api/tts', 'Offline speech-to-text and text-to-speech.'),
])

# ── 5. Gemma Integration ──
h1('5. Gemma Integration Details')
h2('5.1 How Gemma Is Used')
body('Gemma is the reasoning core of अधिKaar. It runs entirely locally through Ollama (gemma4:e4b, with an '
     'automatic gemma3:4b fallback if the primary model is unavailable or crashes on the host). The model is '
     'NOT fine-tuned; instead it is grounded with RAG and carefully prompt-engineered, so its output is anchored '
     'to retrieved official-law text rather than parametric memory. EmbeddingGemma powers the semantic retrieval '
     'layer where enabled.')
h2('5.2 Generation Controls')
bullet('num_ctx (context window) is set to 8192 so a long retrieved-law prompt and a long, detailed answer both fit.')
bullet('num_predict is uncapped (-1) so comprehensive answers finish naturally instead of truncating mid-sentence.')
bullet('Grammar-constrained JSON (Ollama format=) is used for structured outputs (verified analysis, rights card) to guarantee parseable results.')
bullet('A GPU-crash guard automatically retries in CPU mode; MKLDNN is disabled for OCR to avoid a PaddlePaddle oneDNN crash on some CPUs.')
h2('5.3 Value Gemma Brings')
body('Gemma turns messy, real-world input - plain-language questions, code-mixed Hinglish, even garbled OCR of a '
     'scanned FIR - into grounded, citizen-readable legal guidance. Critically, its small on-device footprint is '
     'exactly what makes the offline, private-by-design promise achievable: a capable model that fits on a '
     'citizen\'s own machine, so sensitive legal details never travel to a server.')

# ── 6. Legal Knowledge Base & RAG ──
h1('6. Legal Knowledge Base & Retrieval')
h2('6.1 The Corpus')
body('Answers are grounded on 6,845 reviewed chunks of official Indian law, ingested into a dedicated ChromaDB '
     'collection. Each chunk carries metadata including the act name, section, jurisdiction, effective date and '
     'the official source URL, so citations resolve to real government sources.')
kv_table([
    ('Statutes covered', 'Bharatiya Nyaya Sanhita 2023, Bharatiya Nagarik Suraksha Sanhita 2023, Bharatiya Sakshya Adhiniyam 2023, Constitution of India, Consumer Protection Act 2019, Code on Wages 2019, Delhi Rent Control Act 1958, Legal Services Authorities Act, NALSA regulations, RTI Act 2005 (optional).'),
    ('Curated datasets', 'IPC↔BNS mappings (216), CrPC↔BNSS mappings (80), document-format templates (45), evidence checklists (20), case-study Q&A (22), rights knowledge, PAN-India legal-aid directory.'),
    ('Retrieval', 'Hybrid: multilingual query expansion + keyword boosters + ChromaDB semantic retrieval, with metadata preserved so source URLs survive to the citation.'),
])
h2('6.2 Multilingual Retrieval')
body('User queries in any of the 11 languages are expanded through a synonym map and, where needed, translated to '
     'English search keywords, so a Hinglish phrase like "mera malik ne salary nahi di" reaches the correct '
     'unpaid-wages knowledge even though the corpus is indexed in English.')

# ── 7. Verified Answer Pipeline ──
h1('7. Verified Answer Pipeline (Law & Next Steps)')
body('The flagship "Law & Next Steps" feature does not simply trust the model. It runs a draft → guard → verify → '
     'repair pipeline adapted from claim-level verification research:')
num('DRAFT - the model reads the retrieved statute excerpts and produces the six panels plus a list of claims, each citing the chunk IDs it relied on.')
num('DETERMINISTIC GUARD - a regex hallucination check rejects any claim that names a section number absent from its cited excerpts, before any further model call.')
num('VERIFY - surviving claims are batch-checked against only their cited excerpts; unsupported claims are flagged.')
num('REPAIR - unsupported or unverified claims are excluded from the law, rights and sources panels; the "Sources" panel lists only the official URLs of chunks cited by verified claims.')
body('The result: every citation shown to the user is either backed by retrieved official law or explicitly marked '
     'unverified - directly attacking AI hallucination in a domain where a wrong citation is dangerous.')

# ── 8. Voice, OCR & Accessibility ──
h1('8. Voice, OCR & Accessibility')
h2('8.1 Voice')
bullet('Speech-to-text: faster-whisper transcribes recorded audio on-device, with VAD filtering and hallucination guards (silence no longer becomes phantom "Thank you" text).')
bullet('Text-to-speech: MMS-TTS synthesises audio for all 11 languages locally, with a browser Web Speech API fallback.')
bullet('Karaoke read-aloud: each word is highlighted as it is spoken, aiding low-literacy users.')
bullet('Voice requires a secure context (localhost or HTTPS); a free Cloudflare Tunnel provides HTTPS for remote demos while keeping inference local.')
h2('8.2 Document OCR')
body('Uploaded documents are read on-device. For scanned, often bilingual Indian legal papers, the OCR model is '
     'chosen by the document\'s script rather than the UI language - defaulting to the Devanagari recogniser, '
     'which reads Devanagari and Latin together, so English structured fields (FIR number, sections, names, '
     'vehicle numbers) extract cleanly. A pdf.js text-layer path and Tesseract.js provide in-browser fallbacks.')

# ── 9. Privacy & Security ──
h1('9. Privacy & Security by Design')
bullet('100% on-device: the LLM, embeddings, RAG store, OCR and speech all run locally; the demo works with Wi-Fi disabled.')
bullet('Loopback-only API (127.0.0.1) - the service is not exposed to the network by default.')
bullet('Uploaded documents and transcripts are processed per-session and are not persisted beyond the session by default.')
bullet('Safety boundaries: official sources and effective dates outrank model memory; unsupported legal claims are removed or marked; the assistant provides legal information, not legal advice, and routes to NALSA (15100) for professional help.')

# ── 10. Innovation & Uniqueness ──
h1('10. Innovation & Uniqueness')
body('Most legal tools pick one or two of the following; अधिKaar combines all of them in a single citizen-facing app:')
bullet('Genuinely offline and private - no cloud dependency, so sensitive legal details never leave the device.')
bullet('Claim-level verification - each statement checked against retrieved official law, unsupported claims dropped, real source links shown.')
bullet('Current with the 2023-2024 legal transition - built-in IPC↔BNS and CrPC↔BNSS converters and a BNS/BNSS-first knowledge base.')
bullet('Truly multilingual and multimodal - 11 languages, voice in and voice out, and on-device document OCR with plain-language explanation.')
bullet('Action-oriented - not just answers, but ready-to-file documents, evidence checklists, and PAN-India legal-aid contacts.')

# ── 11. Setup & Deployment ──
h1('11. Setup & Deployment')
h2('11.1 Prerequisites')
bullet('Python 3.11/3.12, Ollama, and the Gemma models (ollama pull gemma4:e4b).')
bullet('Python dependencies from requirements.txt (Flask, ChromaDB, transformers, torch, faster-whisper, paddleocr, etc.).')
h2('11.2 Run')
num('Build the RAG knowledge base once: python rag_setup.py (the official-law corpus build is the slow, one-time step; individual collections rebuild with --only).')
num('Start the service: python app.py - the app is served at http://localhost:5000.')
num('First voice/OCR use downloads the respective models once, then runs fully offline.')
h2('11.3 Distribution')
body('Because the design is local-first, distribution means shipping the package so each user runs it on their own '
     'machine - via a one-command installer or Docker Compose (app + Ollama). For live demos, a free Cloudflare '
     'Tunnel exposes an HTTPS URL while all inference stays on the host machine.')

# ── 12. Limitations & Future Work ──
h1('12. Limitations & Future Work')
bullet('OCR of handwritten or low-quality bilingual scans is imperfect; PP-OCRv6 does not yet include a Devanagari model, so Hindi text uses PP-OCRv5.')
bullet('The corpus is focused on high-impact central acts and demo domains; state-specific laws and more acts can be added via the same ingestion pipeline.')
bullet('Legal content is auto-generated from official sources and curated mappings; a human legal-accuracy review is recommended before production use.')
bullet('Future work: broader corpus, packaged desktop/mobile distribution, per-section deep-link citations, and expanded regional-language coverage.')

# ── 13. Disclaimer ──
h1('13. Disclaimer')
p = doc.add_paragraph()
r = p.add_run('अधिKaar provides legal information, not legal advice. For complex matters, users should consult a '
              'qualified lawyer or call NALSA on 15100. It is an assistive tool that helps citizens understand '
              'their rights and options; it does not replace professional legal counsel.')
r.italic = True; r.font.color.rgb = GREY

out = r'C:\developer\tools\AdhiKaar\docs\AdhiKaar_Documentation.docx'
import os
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print('Saved:', out)
print('Paragraphs:', len(doc.paragraphs))
